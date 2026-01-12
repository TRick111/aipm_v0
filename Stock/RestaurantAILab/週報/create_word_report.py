# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Wordドキュメントを作成
doc = Document()

# タイトルスタイル設定
title = doc.add_heading('Five Arrows 週報作成基礎資料', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# メタ情報
doc.add_paragraph('対象期間: 2025年10月18日〜24日 (2025-W42)')
doc.add_paragraph('作成日: 2025年12月31日')
doc.add_paragraph('データ処理: UTC→JST変換済み、営業日基準')
doc.add_paragraph()

# セクション1: 売上分析
doc.add_heading('1. 売上分析', level=1)

doc.add_heading('1.1 週次推移分析', level=2)
doc.add_paragraph('対象週 (2025-W42) サマリー')
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '指標'
cells[1].text = '実績値'
data = [('売上', '¥624,050'), ('客数', '105人'), ('組数', '49組'), ('客単価', '¥5,943')]
for i, (k, v) in enumerate(data, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph('【前週比・前月比・前年比の比較】')
doc.add_paragraph('・売上: 前週比 -19.5%, 前月比 +84.2%, 前年比 -45.1%')
doc.add_paragraph('・客数: 前週比 -37.1%, 前月比 +38.2%, 前年比 -58.2%')
doc.add_paragraph('・客単価: 前週比 +28.1%')

doc.add_paragraph()
doc.add_paragraph('【インサイト】')
doc.add_paragraph('✅ ポジティブ: 客単価の大幅上昇（+28.1%）、オリジナルカクテル人気')
doc.add_paragraph('⚠️ 懸念: 客数の大幅減少（前週比-37.1%、前年比-58.2%）')
doc.add_paragraph('💡 推奨: 集客施策の強化、SNSプロモーション、予約促進キャンペーン')

doc.add_heading('1.2 曜日別分析（要因分解付き）', level=2)
doc.add_paragraph('【深掘り対象曜日】')
doc.add_paragraph('・客数要因 最大プラス: 月曜日 (+¥28,504)')
doc.add_paragraph('・客数要因 最大マイナス: 火曜日 (-¥216,781)')
doc.add_paragraph('・客単価要因 最大プラス: 金曜日 (+¥90,512)')

doc.add_paragraph()
doc.add_paragraph('【インサイト】')
doc.add_paragraph('✅ 全曜日で客単価要因がプラス')
doc.add_paragraph('⚠️ 火曜日の客数激減（39人→15人、-62%）')
doc.add_paragraph('💡 火曜日限定プロモーション、金曜日予約促進が必要')

# セクション2: 商品・カテゴリ分析
doc.add_heading('2. 商品・カテゴリ分析', level=1)

doc.add_heading('2.1 カテゴリ別分析', level=2)
doc.add_paragraph('【カテゴリ別ランキング TOP5】')
doc.add_paragraph('1. コース&セット: ¥138,000 (23.4%, 前週比-9.1pt)')
doc.add_paragraph('2. オリジナルカクテル: ¥74,700 (12.7%, 前週比+4.5pt)')
doc.add_paragraph('3. ジャパニーズカクテル: ¥54,600 (9.3%, 前週比+1.9pt)')
doc.add_paragraph('4. ウイスキー（ハイボール等）: ¥44,200 (7.5%, 前週比-5.1pt)')
doc.add_paragraph('5. その他: ¥39,400 (6.7%, 前週比-1.0pt)')

doc.add_heading('2.2 商品別分析', level=2)
doc.add_paragraph('【商品別ランキング TOP5】')
doc.add_paragraph('1. 6500円飲み放題付: ¥71,500 (12.12%, 前週比-12.97pt)')
doc.add_paragraph('2. その他CLPコース: ¥66,500 (11.27%, 前週比+3.82pt)')
doc.add_paragraph('3. tablecharge: ¥39,000 (6.61%)')
doc.add_paragraph('4. ハウスハイボール: ¥33,800 (5.73%, 前週比+2.79pt)')
doc.add_paragraph('5. ガージェリー: ¥24,700 (4.19%, 前週比+2.11pt)')

# セクション3: 口コミ分析
doc.add_heading('3. 口コミ分析', level=1)
doc.add_paragraph('対象週の新規投稿数: 0件')
doc.add_paragraph('⚠️ 口コミ投稿がなく、顧客エンゲージメントの機会損失')
doc.add_paragraph('💡 口コミ投稿キャンペーン、SNS情報発信の強化が必要')

# セクション4: アクションプラン
doc.add_heading('4. 総合的な考察とアクションプラン', level=1)

doc.add_heading('4.1 現状総括', level=2)
doc.add_paragraph('【強み】')
doc.add_paragraph('・客単価の上昇（+28.1%）')
doc.add_paragraph('・オリジナルカクテルの人気')
doc.add_paragraph('・深夜売上の安定（23.1%）')

doc.add_paragraph()
doc.add_paragraph('【課題】')
doc.add_paragraph('・客数の大幅減少（最優先課題）')
doc.add_paragraph('・火曜・金曜の集客不振')
doc.add_paragraph('・コース予約の変動')

doc.add_heading('4.2 アクションプラン', level=2)
doc.add_paragraph('【短期施策（1-3ヶ月）】')
doc.add_paragraph('・火曜日プロモーション（ハッピーアワー）')
doc.add_paragraph('・金曜日予約促進（予約特典）')
doc.add_paragraph('・口コミ投稿キャンペーン')
doc.add_paragraph('・SNS投稿強化（週3回）')

doc.add_paragraph()
doc.add_paragraph('【中期施策（3-6ヶ月）】')
doc.add_paragraph('・予約管理システム導入')
doc.add_paragraph('・LINEリピート施策')
doc.add_paragraph('・メニューリニューアル')

doc.add_heading('4.3 KPI設定', level=2)
table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Table Grid'
headers = ['KPI', '現状値', '目標値', '改善率']
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
kpi_data = [
    ('週間売上', '¥624,050', '¥750,000', '+20%'),
    ('週間客数', '105人', '130人', '+24%'),
    ('客単価', '¥5,943', '¥5,800', '-2%'),
    ('火曜日客数', '15人', '25人', '+67%'),
    ('金曜日客数', '15人', '28人', '+87%')
]
for i, row_data in enumerate(kpi_data, 1):
    for j, val in enumerate(row_data):
        table2.rows[i].cells[j].text = val

# 保存
output_path = 'Stock/RestaurantAILab/週報/2_output_2025w42/週報作成基礎資料.docx'
doc.save(output_path)
print(f'Wordファイルを作成しました: {output_path}')



