import os
import sys
import argparse
from pathlib import Path

# 標準出力のエンコーディングを設定
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

print("=" * 80)
print("Word文書（.docx）生成スクリプト")
print("=" * 80)

# コマンドライン引数の解析
parser = argparse.ArgumentParser(description='MarkdownファイルをWord文書(.docx)に変換')
parser.add_argument('-d', '--directory', type=str,
                    default=r'Stock\RestaurantAILab\週報\2_output_2025w42',
                    help='対象ディレクトリのパス (デフォルト: Stock\\RestaurantAILab\\週報\\2_output_2025w42)')
args = parser.parse_args()

# 必要なライブラリのインストール確認
try:
    import pypandoc
    print("✓ pypandoc がインストールされています")
except ImportError:
    print("pypandoc をインストール中...")
    os.system("pip install pypandoc")
    import pypandoc

# pandocのインストール確認
try:
    pypandoc.get_pandoc_version()
    print(f"✓ pandoc バージョン: {pypandoc.get_pandoc_version()}")
except OSError:
    print("\npandoc本体をインストール中...")
    print("これには数分かかる場合があります...")
    try:
        pypandoc.download_pandoc()
        print("✓ pandoc のインストールが完了しました")
    except Exception as e:
        print(f"✗ pandoc のインストールに失敗しました: {e}")
        print("\n手動でインストールする場合:")
        print("1. https://pandoc.org/installing.html にアクセス")
        print("2. Windows版をダウンロードしてインストール")
        sys.exit(1)

# ファイルパスの設定
base_dir = Path(args.directory)
if not base_dir.exists():
    print(f"\n✗ エラー: 指定されたディレクトリが存在しません: {base_dir}")
    sys.exit(1)

md_file = base_dir / '週報作成基礎資料.md'
if not md_file.exists():
    print(f"\n✗ エラー: Markdownファイルが存在しません: {md_file}")
    sys.exit(1)

docx_file = base_dir / '週報作成基礎資料.docx'

print(f"\n読み込み: {md_file}")
print(f"出力先: {docx_file}")

# Markdownファイルを読み込み
with open(md_file, 'r', encoding='utf-8') as f:
    md_content = f.read()

print("\nWord文書を生成中...")

try:
    # pandocのオプション設定
    extra_args = [
        '--standalone',
        '--toc',  # 目次を追加
        '--toc-depth=3',  # 目次の深さ
        '--reference-doc=',  # カスタムテンプレートを使用する場合
        f'--resource-path={base_dir}',  # 画像ファイルのパス
    ]

    # Markdownから.docxに変換
    output = pypandoc.convert_file(
        str(md_file),
        'docx',
        outputfile=str(docx_file),
        extra_args=[
            '--standalone',
            '--toc',
            '--toc-depth=3',
            f'--resource-path={base_dir}',
        ]
    )

    print(f"\n✓ Word文書の生成が完了しました: {docx_file}")

    # ファイルサイズを表示
    file_size = os.path.getsize(docx_file)
    file_size_mb = file_size / (1024 * 1024)
    print(f"  ファイルサイズ: {file_size_mb:.2f} MB")

    print("\n📌 次のステップ:")
    print("1. Googleドライブにアクセス")
    print("2. 「新規」→「ファイルのアップロード」を選択")
    print(f"3. {docx_file} をアップロード")
    print("4. アップロード後、ファイルをダブルクリックしてGoogleドキュメントで開く")
    print("5. 必要に応じて「Googleドキュメント形式で保存」を選択")

except Exception as e:
    print(f"\n✗ Word文書生成中にエラーが発生しました: {e}")
    print(f"  エラータイプ: {type(e).__name__}")

    # 代替案を提案
    print("\n【代替案】python-docxを使用した生成を試みます...")

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        print("✓ python-docx を使用します")
    except ImportError:
        print("python-docx をインストール中...")
        os.system("pip install python-docx")
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

    # セルに背景色を設定する関数
    def set_cell_background(cell, fill_color):
        """セルの背景色を設定"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def set_cell_border(cell):
        """セルに罫線を設定"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # 罫線の設定
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 罫線の太さ
            border.set(qn('w:color'), '000000')  # 黒
            tcBorders.append(border)

        tcPr.append(tcBorders)

    # 新しいWord文書を作成
    doc = Document()

    # タイトルの追加
    title = doc.add_heading('週報作成基礎資料 - Five Arrows', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Markdownの各行を処理
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 見出し
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # 画像
        elif line.startswith('!['):
            match = re.search(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_relative_path = match.group(1).lstrip('./')
                img_path = base_dir / img_relative_path
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as img_error:
                        doc.add_paragraph(f"[画像: {img_relative_path}] (読み込みエラー: {img_error})")
                else:
                    doc.add_paragraph(f"[画像が見つかりません: {img_relative_path}]")

        # 水平線
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 80)

        # テーブル（簡易的な処理）
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # テーブルの開始
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1

            if len(table_lines) > 2:
                # ヘッダーとデータを分離
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                data_rows = []
                for row in table_lines[2:]:  # 2行目は区切り線なのでスキップ
                    cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                    if cells:
                        data_rows.append(cells)

                # テーブルを作成
                if headers and data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                    table.style = 'Table Grid'

                    # ヘッダー行の設定
                    for j, header in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = header
                        # ヘッダーに背景色を設定（濃い青）
                        set_cell_background(cell, '4472C4')
                        # ヘッダーのテキストを白色・太字に
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        set_cell_border(cell)

                    # データ行の設定
                    for i_row, row_data in enumerate(data_rows):
                        for j, cell_data in enumerate(row_data):
                            if j < len(headers):
                                cell = table.rows[i_row + 1].cells[j]
                                cell.text = cell_data

                                # 罫線を設定
                                set_cell_border(cell)

                                # 重要なデータを強調（太字を含む、マイナス値、%表示など）
                                if '**' in cell_data or '-' in cell_data or '%' in cell_data:
                                    # ** を削除して太字に
                                    clean_text = cell_data.replace('**', '')
                                    cell.text = clean_text

                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True

                                            # マイナス値は赤色
                                            if '-' in clean_text and '%' in clean_text:
                                                run.font.color.rgb = RGBColor(192, 0, 0)
                                            # プラス値は緑色
                                            elif '+' in clean_text and '%' in clean_text:
                                                run.font.color.rgb = RGBColor(0, 128, 0)

                                # 偶数行に薄い背景色を設定（ゼブラストライプ）
                                if i_row % 2 == 0:
                                    set_cell_background(cell, 'F2F2F2')

        # 箇条書き
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')

        # 通常の段落
        elif line:
            # 太字や強調の処理（簡易版）
            p = doc.add_paragraph(line)

        i += 1

    # 保存
    doc.save(str(docx_file))

    file_size = os.path.getsize(docx_file)
    file_size_mb = file_size / (1024 * 1024)
    print(f"\n✓ python-docxでWord文書を生成しました: {docx_file}")
    print(f"  ファイルサイズ: {file_size_mb:.2f} MB")

    print("\n📌 次のステップ:")
    print("1. Googleドライブにアクセス")
    print("2. 「新規」→「ファイルのアップロード」を選択")
    print(f"3. {docx_file} をアップロード")
    print("4. アップロード後、ファイルをダブルクリックしてGoogleドキュメントで開く")

print("\n" + "=" * 80)
print("処理完了")
print("=" * 80)
